"""
ISO 9001:2015 kompatibilis munkalap sablon generáló
Professzionális design, teljes nyomonkövethetőség
"""

import os
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("python-docx not installed. Please install it with: pip install python-docx")
    exit(1)


def set_cell_bg(cell, color_hex):
    """Set background color of a table cell."""
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_border(cell, color_hex="000000", width_pt=4):
    """Set cell border."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{width_pt}" w:space="0" w:color="{color_hex}"/>'
        f'<w:left w:val="single" w:sz="{width_pt}" w:space="0" w:color="{color_hex}"/>'
        f'<w:bottom w:val="single" w:sz="{width_pt}" w:space="0" w:color="{color_hex}"/>'
        f'<w:right w:val="single" w:sz="{width_pt}" w:space="0" w:color="{color_hex}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def configure_styles(doc):
    """Setup custom styles for ISO 9001 compliant professional look."""
    # Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)

    # Heading 1 - Document Title
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(20)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x00, 0x37, 0x6F)  # ISO Blue
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(12)

    # Heading 2 - Section Headers
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x00, 0x37, 0x6F)  # ISO Blue
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)

    # Table Header style
    styles = doc.styles
    if 'Table Header' not in styles:
        th_style = styles.add_style('Table Header', 1)
        th_font = th_style.font
        th_font.name = 'Calibri'
        th_font.size = Pt(10)
        th_font.bold = True
        th_font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)  # White
        th_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def create_iso9001_template(output_path):
    """Create ISO 9001:2015 compliant worksheet template."""
    doc = Document()
    
    # Page setup - ISO standard margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)

    configure_styles(doc)

    # ========================================================================
    # HEADER SECTION - ISO 9001 Document Control
    # ========================================================================
    header_table = doc.add_table(rows=4, cols=2)
    header_table.width = Inches(7.5)
    
    # Header row 1: Document Title
    header_table.cell(0, 0).merge(header_table.cell(0, 1))
    title_cell = header_table.cell(0, 0)
    title_para = title_cell.paragraphs[0]
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("MUNKALAP / WORK ORDER")
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x00, 0x37, 0x6F)
    set_cell_bg(title_cell, "E7F3FF")
    set_cell_border(title_cell, "00376F", 8)
    
    # Header row 2: Document Number and Version
    doc_num_cell = header_table.cell(1, 0)
    doc_num_cell.width = Cm(4)
    doc_num_para = doc_num_cell.paragraphs[0]
    doc_num_para.add_run("Dokumentum szám / Document No.:").bold = True
    doc_num_para.add_run("\nWS-${worksheet.id}")
    set_cell_bg(doc_num_cell, "F2F4F6")
    set_cell_border(doc_num_cell)
    
    version_cell = header_table.cell(1, 1)
    version_cell.width = Cm(4)
    version_para = version_cell.paragraphs[0]
    version_para.add_run("Verzió / Version:").bold = True
    version_para.add_run("\n1.0")
    set_cell_bg(version_cell, "F2F4F6")
    set_cell_border(version_cell)
    
    # Header row 3: Generation Info
    gen_cell = header_table.cell(2, 0)
    gen_para = gen_cell.paragraphs[0]
    gen_para.add_run("Generálva / Generated:").bold = True
    gen_para.add_run("\n${generated_at}")
    set_cell_bg(gen_cell, "F2F4F6")
    set_cell_border(gen_cell)
    
    generator_cell = header_table.cell(2, 1)
    generator_para = generator_cell.paragraphs[0]
    generator_para.add_run("Generálta / Generated by:").bold = True
    generator_para.add_run("\n${generated_by}")
    set_cell_bg(generator_cell, "F2F4F6")
    set_cell_border(generator_cell)
    
    # Header row 4: ISO 9001 Compliance Notice
    iso_cell = header_table.cell(3, 0).merge(header_table.cell(3, 1))
    iso_para = iso_cell.paragraphs[0]
    iso_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    iso_run = iso_para.add_run("ISO 9001:2015 nyomonkövethetőségi dokumentum / Traceability Document")
    iso_run.font.size = Pt(8)
    iso_run.font.italic = True
    iso_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    set_cell_bg(iso_cell, "FAFAFA")
    set_cell_border(iso_cell)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ========================================================================
    # SECTION 1: TRACEABILITY DATA (ISO 9001 Requirement)
    # ========================================================================
    doc.add_heading('1. NYOMONKÖVETETHETŐSÉGI ADATOK / TRACEABILITY DATA', level=2)
    
    trace_data = [
        ("Státusz / Status", "${worksheet.status}"),
        ("Létrehozva / Created at", "${worksheet.created_at}"),
        ("Hiba kezdete / Breakdown start", "${worksheet.breakdown_time}"),
        ("Javítás befejezése / Repair finished", "${worksheet.repair_finished_time}"),
        ("Leállási idő / Downtime", "${worksheet.downtime}"),
        ("Felelős technikus / Responsible technician", "${worksheet.assigned_username}"),
    ]
    
    trace_table = doc.add_table(rows=len(trace_data), cols=2)
    trace_table.style = 'Table Grid'
    
    for i, (label, value) in enumerate(trace_data):
        row = trace_table.rows[i]
        # Label cell
        label_cell = row.cells[0]
        label_cell.width = Cm(6)
        label_para = label_cell.paragraphs[0]
        label_run = label_para.add_run(label)
        label_run.font.bold = True
        label_run.font.size = Pt(10)
        set_cell_bg(label_cell, "E7F3FF")
        set_cell_border(label_cell)
        
        # Value cell
        value_cell = row.cells[1]
        value_cell.width = Cm(8)
        value_para = value_cell.paragraphs[0]
        value_para.add_run(value)
        value_para.paragraph_format.space_after = Pt(0)
        set_cell_border(value_cell)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ========================================================================
    # SECTION 2: EQUIPMENT INFORMATION
    # ========================================================================
    doc.add_heading('2. BERENDEZÉS ADATAI / EQUIPMENT INFORMATION', level=2)
    
    equip_data = [
        ("Gép neve / Machine name", "${machine.name}"),
        ("Gyári szám / Serial number", "${machine.serial_number}"),
        ("Gyártó / Modell / Manufacturer / Model", "${machine.manufacturer_model}"),
        ("Termelési vonal / Production line", "${production_line.name}"),
    ]
    
    equip_table = doc.add_table(rows=len(equip_data), cols=2)
    equip_table.style = 'Table Grid'
    
    for i, (label, value) in enumerate(equip_data):
        row = equip_table.rows[i]
        label_cell = row.cells[0]
        label_cell.width = Cm(6)
        label_para = label_cell.paragraphs[0]
        label_run = label_para.add_run(label)
        label_run.font.bold = True
        label_run.font.size = Pt(10)
        set_cell_bg(label_cell, "E7F3FF")
        set_cell_border(label_cell)
        
        value_cell = row.cells[1]
        value_cell.width = Cm(8)
        value_para = value_cell.paragraphs[0]
        value_para.add_run(value)
        set_cell_border(value_cell)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ========================================================================
    # SECTION 3: FAULT DESCRIPTION / NOTES
    # ========================================================================
    doc.add_heading('3. HIBA LEÍRÁSA / MEGJEGYZÉSEK / FAULT DESCRIPTION / NOTES', level=2)
    
    notes_para = doc.add_paragraph("${worksheet.notes}")
    notes_para.paragraph_format.left_indent = Cm(0.5)
    notes_para.paragraph_format.space_after = Pt(12)
    
    # Add border box around notes (using a table)
    notes_table = doc.add_table(rows=1, cols=1)
    notes_table.width = Inches(7.5)
    notes_cell = notes_table.cell(0, 0)
    notes_cell.text = "${worksheet.notes}"
    set_cell_border(notes_cell, "00376F", 6)
    notes_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ========================================================================
    # SECTION 4: USED PARTS (ISO 9001 Material Traceability)
    # ========================================================================
    doc.add_heading('4. FELHASZNÁLT ALKATRÉSZEK / USED PARTS', level=2)
    
    # Placeholder table - will be populated by the application
    parts_table = doc.add_table(rows=2, cols=5)
    parts_table.style = 'Table Grid'
    
    # Header row
    headers = [
        "SKU",
        "Megnevezés / Name",
        "Mennyiség / Qty",
        "Egységköltség / Unit cost",
        "Összköltség / Total cost"
    ]
    
    header_row = parts_table.rows[0]
    for i, header_text in enumerate(headers):
        header_cell = header_row.cells[i]
        header_para = header_cell.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_para.add_run(header_text)
        header_run.font.bold = True
        header_run.font.size = Pt(10)
        set_cell_bg(header_cell, "00376F")
        set_cell_border(header_cell)
    
    # Placeholder row
    placeholder_row = parts_table.rows[1]
    placeholder_cell = placeholder_row.cells[0].merge(placeholder_row.cells[4])
    placeholder_para = placeholder_cell.paragraphs[0]
    placeholder_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    placeholder_run = placeholder_para.add_run("[Az alkatrész táblázatot az alkalmazás automatikusan kitölti / Parts table will be auto-populated by application]")
    placeholder_run.font.italic = True
    placeholder_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    set_cell_border(placeholder_cell)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ========================================================================
    # SECTION 5: APPROVAL & SIGNATURES (ISO 9001 Requirement)
    # ========================================================================
    doc.add_heading('5. ALÁÍRÁSOK ÉS JÓVÁHAGYÁS / SIGNATURES & APPROVAL', level=2)
    
    sig_table = doc.add_table(rows=3, cols=3)
    sig_table.width = Inches(7.5)
    
    # Header row
    sig_headers = [
        "Elvégezte / Performed by",
        "Ellenőrizte / Checked by",
        "Jóváhagyta / Approved by"
    ]
    
    sig_header_row = sig_table.rows[0]
    for i, header_text in enumerate(sig_headers):
        header_cell = sig_header_row.cells[i]
        header_cell.width = Cm(4.5)
        header_para = header_cell.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_para.add_run(header_text)
        header_run.font.bold = True
        header_run.font.size = Pt(10)
        set_cell_bg(header_cell, "00376F")
        set_cell_border(header_cell)
    
    # Signature space row
    sig_space_row = sig_table.rows[1]
    for cell in sig_space_row.cells:
        cell.height = Cm(3)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        para = cell.paragraphs[0]
        para.add_run("\n\n_________________")
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_border(cell)
    
    # Date row
    date_row = sig_table.rows[2]
    for cell in date_row.cells:
        para = cell.paragraphs[0]
        para.add_run("Dátum / Date: _______________")
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(0)
        set_cell_border(cell)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ========================================================================
    # FOOTER NOTE
    # ========================================================================
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        "Ez a dokumentum automatikusan generált, ISO 9001:2015 nyomonkövethetőségi követelményeknek megfelelően.\n"
        "This document is automatically generated, compliant with ISO 9001:2015 traceability requirements."
    )
    footer_run.font.size = Pt(8)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Save
    doc.save(output_path)
    print(f"ISO 9001 template generated successfully at: {output_path}")


if __name__ == "__main__":
    # Output to templates directory
    templates_dir = Path(__file__).parent.parent / "templates"
    templates_dir.mkdir(parents=True, exist_ok=True)
    output_file = templates_dir / "worksheet_template.docx"
    
    create_iso9001_template(output_file)
    print(f"\nTemplate saved to: {output_file}")
    print("This template is ISO 9001:2015 compliant and ready for document control.")

