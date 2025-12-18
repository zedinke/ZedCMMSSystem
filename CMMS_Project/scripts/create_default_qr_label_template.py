"""
Script to create default QR label template
Creates a professional DOCX template for QR code labels
"""

import sys
from pathlib import Path

# Add project root to path
PROJECT_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from config.app_config import TEMPLATES_DIR

def create_default_qr_label_template():
    """Create a default QR label template DOCX file"""
    
    # Ensure templates directory exists
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
    
    # Template file path
    template_path = TEMPLATES_DIR / "default_qr_label_template.docx"
    
    # Create new document
    doc = Document()
    
    # Set page margins (A4: 21cm x 29.7cm)
    # Small margins for label printing
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
    
    # Add title/header (optional - will be replaced by actual labels)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("QR Címke Sablon / QR Label Template")
    title_run.font.size = Pt(10)
    title_run.font.color.rgb = RGBColor(100, 100, 100)
    title_run.italic = True
    
    # Add spacing
    doc.add_paragraph()
    
    # Add a sample label structure as reference
    # This will show the layout but will be replaced when generating actual labels
    sample_table = doc.add_table(rows=1, cols=2)
    sample_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Left cell - text data area
    left_cell = sample_table.cell(0, 0)
    left_cell.width = Cm(6)  # 60% of 10cm label
    
    # Right cell - QR code area
    right_cell = sample_table.cell(0, 1)
    right_cell.width = Cm(4)  # 40% of 10cm label
    
    # Add sample text to left cell
    p_left = left_cell.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # SKU
    run = p_left.add_run("SKU: [SKU]\n")
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Name
    run = p_left.add_run("[Név / Name]\n")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Machine info
    run = p_left.add_run("Szülő: [Szülő gép]\n")
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    run = p_left.add_run("Alárendelt: [Alárendelt gép]\n")
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    # Creator
    run = p_left.add_run("Készítette:\n[Készítő]\n")
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    # Date
    run = p_left.add_run("Dátum:\n[Dátum]")
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    # Add placeholder text to right cell
    p_right = right_cell.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_right.add_run("[QR Code]\n2.5cm x 2.5cm")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(200, 200, 200)
    run.italic = True
    
    # Save template
    doc.save(str(template_path))
    print(f"✓ Default QR label template created: {template_path}")
    return template_path


if __name__ == "__main__":
    template_path = create_default_qr_label_template()
    print(f"\nTemplate saved to: {template_path}")
    print("You can now set this as the default template in settings.")

