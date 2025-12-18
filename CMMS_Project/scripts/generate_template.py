import os
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("python-docx not installed. Please install it with: pip install python-docx")
    exit(1)

def set_cell_bg(cell, color_hex):
    """Set background color of a table cell."""
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def configure_styles(doc):
    """Setup custom styles for a modern look."""
    # Basic font settings
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(10)

    # Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Segoe UI'
    h1.font.size = Pt(24)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50) # Dark Blue-Grey
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)

    # Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Segoe UI'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)
    
    # Table Header font
    styles = doc.styles
    if 'Table Header' not in styles:
        th_style = styles.add_style('Table Header', 1) # 1 is Paragraph style
        th_font = th_style.font
        th_font.name = 'Segoe UI'
        th_font.size = Pt(10)
        th_font.bold = True
        th_font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) # White

def create_modern_template(output_path):
    doc = Document()
    
    # Page margins (Narrower for more space)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    configure_styles(doc)

    # --- HEADER SECTION (Logo + Info) ---
    # Create a 2-column table for the header
    header_table = doc.add_table(rows=1, cols=2)
    header_table.width = Inches(10) # Full width
    header_table.autofit = False
    
    # Column widths: Logo (Left) gets ~30%, Info (Right) gets ~70%
    # Note: python-docx table column widths are tricky, setting cell widths helps
    cell_logo = header_table.cell(0, 0)
    cell_logo.width = Cm(5)
    cell_info = header_table.cell(0, 1)
    cell_info.width = Cm(11)

    # Logo Placeholder
    p = cell_logo.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[LOGO PLACEHOLDER]")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0xBD, 0xC3, 0xC7) # Light Grey
    # You could add an image here with: run.add_picture('path/to/logo.png', width=Cm(4))
    
    # Document Info
    p_info = cell_info.paragraphs[0]
    p_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_title = p_info.add_run("Munkalap / Work Order\n")
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    
    p_info.add_run(f"No: WS-${{worksheet.id}}\n").bold = True
    p_info.add_run("Dokumentum / Document: WS-${worksheet.id} | Verzió / Version: 1.0\n")
    p_info.add_run("Generálva / Generated: ${generated_at} | Generálta: ${generated_by}")

    doc.add_paragraph().paragraph_format.space_after = Pt(12) # Spacer

    # --- TRACEABILITY SECTION ---
    doc.add_heading('Nyomonkövethetőségi adatok / Traceability', level=2)
    
    data_labels = [
        ("Státusz / Status", "${worksheet.status}"),
        ("Létrehozva / Created at", "${worksheet.created_at}"),
        ("Hiba kezdete / Breakdown start", "${worksheet.breakdown_time}"),
        ("Javítás befejezése / Repair finished", "${worksheet.repair_finished_time}"),
        ("Leállási idő / Downtime", "${worksheet.downtime}"),
        ("Felelős / Responsible", "${worksheet.assigned_username}")
    ]
    
    table = doc.add_table(rows=len(data_labels), cols=2)
    table.style = 'Table Grid'
    
    for i, (label, value) in enumerate(data_labels):
        row = table.rows[i]
        # Label cell
        c1 = row.cells[0]
        c1.text = label
        c1.paragraphs[0].runs[0].font.bold = True
        set_cell_bg(c1, "F2F4F6") # Light Grey BG
        c1.width = Cm(6)
        
        # Value cell
        c2 = row.cells[1]
        c2.text = value
        
    doc.add_paragraph() # Spacer

    # --- EQUIPMENT SECTION ---
    doc.add_heading('Berendezés adatai / Equipment', level=2)
    
    equip_labels = [
        ("Gép / Machine", "${machine.name}"),
        ("Gyári szám / Serial no.", "${machine.serial_number}"),
        ("Gyártó / Modell / Manufacturer / Model", "${machine.manufacturer_model}"),
        ("Termelési vonal / Production line", "${production_line.name}")
    ]
    
    etable = doc.add_table(rows=len(equip_labels), cols=2)
    etable.style = 'Table Grid'
    
    for i, (label, value) in enumerate(equip_labels):
        row = etable.rows[i]
        c1 = row.cells[0]
        c1.text = label
        c1.paragraphs[0].runs[0].font.bold = True
        set_cell_bg(c1, "F2F4F6")
        c1.width = Cm(6)
        
        c2 = row.cells[1]
        c2.text = value

    doc.add_paragraph()

    # --- NOTES SECTION ---
    doc.add_heading('Megjegyzések / Hiba leírás / Notes / Fault description', level=2)
    # Add a bordered box or just text for notes
    p_notes = doc.add_paragraph("${worksheet.notes}")
    p_notes.style = 'Normal'

    doc.add_paragraph()

    # --- USED PARTS SECTION ---
    doc.add_heading('Felhasznált alkatrészek / Used parts', level=2)
    doc.add_paragraph("[Az alkatrész táblázatot az alkalmazás tölti ki / Parts table will be populated by app]")

    doc.add_paragraph()

    # --- SIGNATURES SECTION ---
    doc.add_heading('Aláírások és jóváhagyás / Signatures & approval', level=2)
    
    sig_table = doc.add_table(rows=2, cols=3)
    sig_table.style = 'Table Grid'
    
    headers = ["Elvégezte / Performed by", "Ellenőrizte / Checked by", "Jóváhagyta / Approved by"]
    
    # Header Row
    hdr_row = sig_table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = text
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_bg(cell, "F2F4F6")
        
    # Signature Space Row
    sig_row = sig_table.rows[1]
    for cell in sig_row.cells:
        cell.text = "\n\n\n\n" # Space for signature

    # Save
    doc.save(output_path)
    print(f"Template generated successfully at: {output_path}")

if __name__ == "__main__":
    output_file = "generated_worksheet_template.docx"
    # If run from scripts dir, place it in parent root or wherever convenient
    # But user wants to replace 'templates/worksheet_template.docx' eventually
    
    # For now, let's output to the current directory so we can find it
    create_modern_template(output_file)
