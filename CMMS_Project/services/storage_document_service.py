"""
Storage Document Service
Generates DOCX documents for storage operations (receipt, transfer)
"""

from pathlib import Path
from typing import Optional, Dict
from datetime import datetime
from sqlalchemy.orm import Session, joinedload

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import RGBColor, Inches, Cm, Pt
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False
    WD_ALIGN_PARAGRAPH = None
    RGBColor = None
    Inches = None
    Cm = None
    Pt = None

from database.session_manager import SessionLocal
from database.models import PartLocation, StorageLocation, Part, User, utcnow
from services.settings_service import get_selected_storage_receipt_template, get_selected_storage_transfer_template
from config.app_config import TEMPLATES_DIR
from services.storage_service import get_storage_location_path
import logging

logger = logging.getLogger(__name__)


def _ensure_output_dir() -> Path:
    """Ensure document output directory exists"""
    output_dir = Path.cwd() / "generated_documents"
    output_dir.mkdir(exist_ok=True)
    return output_dir


def _replace_placeholders_in_docx(doc_path: Path, replacements: dict) -> Document:
    """Replace placeholders in DOCX template and return modified document"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not available")
    
    doc = Document(str(doc_path))
    
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        original_text = paragraph.text
        substituted_text = original_text
        for key, value in replacements.items():
            substituted_text = substituted_text.replace(key, str(value) if value is not None else "-")
        
        if original_text != substituted_text:
            # Clear all runs and rebuild
            for _ in range(len(paragraph.runs)):
                run_elem = paragraph.runs[0]._element
                run_elem.getparent().remove(run_elem)
            if substituted_text:
                paragraph.add_run(substituted_text)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Replace in paragraphs within cells
                for paragraph in cell.paragraphs:
                    original_text = paragraph.text
                    substituted_text = original_text
                    for key, value in replacements.items():
                        substituted_text = substituted_text.replace(key, str(value) if value is not None else "-")
                    
                    if original_text != substituted_text:
                        # Clear all runs and rebuild
                        for _ in range(len(paragraph.runs)):
                            run_elem = paragraph.runs[0]._element
                            run_elem.getparent().remove(run_elem)
                        if substituted_text:
                            paragraph.add_run(substituted_text)
    
    return doc


def generate_storage_receipt_document(
    part_location_id: int,
    output_path: Optional[Path] = None,
    session: Session = None
) -> Path:
    """
    Generate storage receipt document (DOCX) when part is assigned to location
    
    Args:
        part_location_id: PartLocation ID
        output_path: Optional output path
        session: Database session
    
    Returns:
        Path: Path to generated DOCX file
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not available")
    
    if session is None:
        session = SessionLocal()
        should_close = True
    else:
        should_close = False
    
    try:
        # Get part location with relationships
        part_location = (
            session.query(PartLocation)
            .options(
                joinedload(PartLocation.part),
                joinedload(PartLocation.storage_location),
                joinedload(PartLocation.assigned_by_user),
            )
            .filter_by(id=part_location_id)
            .first()
        )
        
        if not part_location:
            raise ValueError(f"Part location with id {part_location_id} not found")
        
        # Get template
        template_path = get_selected_storage_receipt_template(session=session)
        if not template_path or not template_path.exists():
            raise ValueError("Storage receipt template not found. Please configure it in settings.")
        
        # Get location path
        location_path = get_storage_location_path(part_location.storage_location_id, session)
        
        # Prepare replacements
        replacements = {
            "{PART_NAME}": part_location.part.name if part_location.part else "-",
            "{PART_SKU}": part_location.part.sku if part_location.part else "-",
            "{QUANTITY}": str(part_location.quantity),
            "{LOCATION_NAME}": part_location.storage_location.name if part_location.storage_location else "-",
            "{LOCATION_PATH}": location_path,
            "{LOCATION_CODE}": part_location.storage_location.code if part_location.storage_location and part_location.storage_location.code else "-",
            "{ASSIGNED_DATE}": part_location.assigned_date.strftime("%Y-%m-%d %H:%M") if part_location.assigned_date else "-",
            "{ASSIGNED_BY}": part_location.assigned_by_user.username if part_location.assigned_by_user else "-",
            "{NOTES}": part_location.notes if part_location.notes else "-",
            "{CURRENT_DATE}": datetime.now().strftime("%Y-%m-%d %H:%M"),
        }
        
        # Generate document
        doc = _replace_placeholders_in_docx(template_path, replacements)
        
        # Save document
        if output_path is None:
            output_dir = _ensure_output_dir()
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = output_dir / f"storage_receipt_{part_location_id}_{timestamp}.docx"
        
        doc.save(str(output_path))
        logger.info(f"Generated storage receipt document: {output_path}")
        return output_path
    
    finally:
        if should_close:
            session.close()


def generate_storage_transfer_document(
    source_part_location_id: int,
    target_location_id: int,
    quantity: int,
    notes: Optional[str] = None,
    output_path: Optional[Path] = None,
    session: Session = None
) -> Path:
    """
    Generate storage transfer document (DOCX) when part is transferred between locations
    
    Args:
        source_part_location_id: Source PartLocation ID
        target_location_id: Target location ID
        quantity: Transfer quantity
        output_path: Optional output path
        session: Database session
    
    Returns:
        Path: Path to generated DOCX file
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not available")
    
    if session is None:
        session = SessionLocal()
        should_close = True
    else:
        should_close = False
    
    try:
        # Get source part location
        source_part_location = (
            session.query(PartLocation)
            .options(
                joinedload(PartLocation.part),
                joinedload(PartLocation.storage_location),
            )
            .filter_by(id=source_part_location_id)
            .first()
        )
        
        if not source_part_location:
            raise ValueError(f"Source part location with id {source_part_location_id} not found")
        
        # Get target location
        target_location = session.query(StorageLocation).filter_by(id=target_location_id).first()
        if not target_location:
            raise ValueError(f"Target location with id {target_location_id} not found")
        
        # Get template
        template_path = get_selected_storage_transfer_template(session=session)
        if not template_path or not template_path.exists():
            raise ValueError("Storage transfer template not found. Please configure it in settings.")
        
        # Get location paths
        source_location_path = get_storage_location_path(source_part_location.storage_location_id, session)
        target_location_path = get_storage_location_path(target_location_id, session)
        
        # Use provided notes, or try to get from target part location if it exists
        transfer_notes = notes
        if not transfer_notes:
            # Try to get notes from target location if part already exists there
            target_part_location = (
                session.query(PartLocation)
                .filter_by(
                    part_id=source_part_location.part_id,
                    storage_location_id=target_location_id
                )
                .first()
            )
            if target_part_location and target_part_location.notes:
                transfer_notes = target_part_location.notes
        
        # Prepare replacements
        replacements = {
            "{PART_NAME}": source_part_location.part.name if source_part_location.part else "-",
            "{PART_SKU}": source_part_location.part.sku if source_part_location.part else "-",
            "{QUANTITY}": str(quantity),
            "{SOURCE_LOCATION_NAME}": source_part_location.storage_location.name if source_part_location.storage_location else "-",
            "{SOURCE_LOCATION_PATH}": source_location_path,
            "{SOURCE_LOCATION_CODE}": source_part_location.storage_location.code if source_part_location.storage_location and source_part_location.storage_location.code else "-",
            "{TARGET_LOCATION_NAME}": target_location.name,
            "{TARGET_LOCATION_PATH}": target_location_path,
            "{TARGET_LOCATION_CODE}": target_location.code if target_location.code else "-",
            "{TRANSFER_DATE}": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "{CURRENT_DATE}": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "{NOTES}": transfer_notes if transfer_notes else "-",
        }
        
        # Generate document
        doc = _replace_placeholders_in_docx(template_path, replacements)
        
        # Save document
        if output_path is None:
            output_dir = _ensure_output_dir()
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = output_dir / f"storage_transfer_{source_part_location_id}_{target_location_id}_{timestamp}.docx"
        
        doc.save(str(output_path))
        logger.info(f"Generated storage transfer document: {output_path}")
        return output_path
    
    finally:
        if should_close:
            session.close()

