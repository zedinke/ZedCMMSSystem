"""
DOCX generálás munkalapokhoz (ISO 9001 nyomonkövetés)
python-docx alapú template rendszerrel.
"""

from pathlib import Path
import json
from typing import Optional
from datetime import datetime

from sqlalchemy.orm import Session, joinedload
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable, KeepTogether, Flowable, PageBreak
from reportlab.graphics.barcode import qr
from reportlab.graphics.shapes import Drawing
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from tempfile import TemporaryDirectory

try:
    from PyPDF2 import PdfReader, PdfWriter
    PYPDF2_AVAILABLE = True
except ImportError:
    try:
        from pypdf import PdfReader, PdfWriter
        PYPDF2_AVAILABLE = True
    except ImportError:
        PYPDF2_AVAILABLE = False

try:
    from docx import Document  # python-docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.shared import RGBColor, Inches, Cm, Pt
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False
    WD_ALIGN_PARAGRAPH = None
    WD_ALIGN_VERTICAL = None
    RGBColor = None
    Inches = None
    Cm = None
    Pt = None

try:
    from docx2pdf import convert as docx_to_pdf_convert
    DOCX2PDF_AVAILABLE = True
except Exception:
    DOCX2PDF_AVAILABLE = False

from database.models import Worksheet, WorksheetPart, WorksheetPDF, Machine, utcnow, PMTask, PMHistory, WorkRequestPDF, PMWorksheetPDF, User, Part, VacationRequest, VacationDocument
from services.worksheet_service import _get_session
from config.app_config import TEMPLATES_DIR
from services.settings_service import get_selected_worksheet_template, get_selected_work_request_template, get_selected_scrapping_template

import logging

logger = logging.getLogger(__name__)


def _ensure_output_dir() -> Path:
    """Ensure PDF output directory exists"""
    output_dir = Path.cwd() / "generated_pdfs"
    output_dir.mkdir(exist_ok=True)
    return output_dir


def _format_dt(dt: Optional[datetime]) -> str:
    """Format datetime for display"""
    if not dt:
        return "-"
    return dt.strftime("%Y-%m-%d %H:%M")


def _replace_placeholders_in_docx(doc_path: Path, replacements: dict) -> Document:
    """Replace placeholders in DOCX template and return modified document"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not available")
    
    doc = Document(str(doc_path))
    
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        original_text = paragraph.text
        substituted_text = original_text
        for key, value in replacements.items():
            substituted_text = substituted_text.replace(key, str(value) if value is not None else "-")
        
        if original_text != substituted_text:
            # Clear all runs and rebuild
            for _ in range(len(paragraph.runs)):
                run_elem = paragraph.runs[0]._element
                run_elem.getparent().remove(run_elem)
            if substituted_text:
                paragraph.add_run(substituted_text)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Replace in paragraphs within cells
                for paragraph in cell.paragraphs:
                    original_text = paragraph.text
                    substituted_text = original_text
                    for key, value in replacements.items():
                        substituted_text = substituted_text.replace(key, str(value) if value is not None else "-")
                    
                    if original_text != substituted_text:
                        # Clear all runs and rebuild
                        for _ in range(len(paragraph.runs)):
                            run_elem = paragraph.runs[0]._element
                            run_elem.getparent().remove(run_elem)
                        if substituted_text:
                            paragraph.add_run(substituted_text)
                
                # Also replace directly in cell text (in case there's no paragraph)
                if cell.text:
                    original_cell_text = cell.text
                    substituted_cell_text = original_cell_text
                    for key, value in replacements.items():
                        substituted_cell_text = substituted_cell_text.replace(key, str(value) if value is not None else "-")
                    
                    if original_cell_text != substituted_cell_text:
                        # Clear all paragraphs and set new text
                        cell.paragraphs.clear()
                        if substituted_cell_text:
                            cell.paragraphs[0].add_run(substituted_cell_text)
    
    return doc


def _populate_parts_table(doc: Document, parts: list):
    """Populate the parts table in the DOCX document with used parts data"""
    if not DOCX_AVAILABLE:
        return
    
    # Find the parts table (look for table with "SKU" in first header cell)
    parts_table = None
    for table in doc.tables:
        if len(table.rows) > 0:
            first_row = table.rows[0]
            if len(first_row.cells) >= 5:
                first_cell_text = first_row.cells[0].text.strip().upper()
                if "SKU" in first_cell_text:
                    parts_table = table
                    break
    
    if not parts_table:
        logger.warning("Parts table not found in template")
        return
    
    # Remove placeholder row(s) - keep only header row
    while len(parts_table.rows) > 1:
        parts_table._element.remove(parts_table.rows[1]._element)
    
    # Add rows for each part
    for part_link in parts:
        part = part_link.part
        if not part:
            continue
        
        row = parts_table.add_row()
        
        # SKU
        row.cells[0].text = part.sku or "-"
        
        # Name
        row.cells[1].text = part.name or "-"
        
        # Quantity
        row.cells[2].text = str(part_link.quantity_used) if part_link.quantity_used else "0"
        
        # Unit cost (from unit_cost_at_time or buy_price)
        unit_cost = part_link.unit_cost_at_time if part_link.unit_cost_at_time else (part.buy_price or 0.0)
        row.cells[3].text = f"{unit_cost:.2f}"
        
        # Total cost
        total_cost = unit_cost * (part_link.quantity_used or 0)
        row.cells[4].text = f"{total_cost:.2f}"
        
        # Style the cells
        for cell in row.cells:
            if cell.paragraphs and len(cell.paragraphs) > 0:
                para = cell.paragraphs[0]
                if len(para.runs) > 0:
                    para.runs[0].font.size = Pt(9)
                else:
                    # If no runs, add one
                    run = para.add_run(cell.text)
                    run.font.size = Pt(9)
                try:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                except:
                    pass  # Ignore if vertical alignment is not supported
    
    # If no parts, add a message row
    if len(parts) == 0:
        row = parts_table.add_row()
        merged_cell = row.cells[0].merge(row.cells[4])
        merged_cell.text = "Nincs felhasznált alkatrész / No parts used"
        if merged_cell.paragraphs and len(merged_cell.paragraphs) > 0:
            merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if len(merged_cell.paragraphs[0].runs) > 0:
                merged_cell.paragraphs[0].runs[0].font.italic = True
                merged_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)


# Removed _docx_to_pdf - we'll keep DOCX format instead of converting to PDF


def generate_worksheet_pdf(worksheet_id: int, generated_by: str = "system", output_path: str | None = None, session: Session = None) -> str:
    """Generate worksheet DOCX using template (no PDF generation)"""
    session, should_close = _get_session(session)
    try:
        ws: Worksheet = (
            session.query(Worksheet)
            .options(
                joinedload(Worksheet.machine).joinedload(Machine.production_line),
                joinedload(Worksheet.parts).joinedload(WorksheetPart.part),
                joinedload(Worksheet.assigned_user),
            )
            .filter(Worksheet.id == worksheet_id)
            .first()
        )
        if not ws:
            raise ValueError("Munkalap nem található")

        if output_path:
            docx_path = Path(output_path)
            if docx_path.suffix.lower() != '.docx':
                docx_path = docx_path.with_suffix('.docx')
            docx_path.parent.mkdir(parents=True, exist_ok=True)
        else:
            output_dir = _ensure_output_dir()
            docx_path = output_dir / f"worksheet_{ws.id}.docx"

        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available for DOCX generation")

        # Use worksheet DOCX template
        template_path = get_selected_worksheet_template(session) or (TEMPLATES_DIR / "worksheet_template.docx")
        if not template_path.exists():
            raise FileNotFoundError(f"Worksheet template not found: {template_path}")

        # Prepare replacements for worksheet template
        replacements = {
            "${worksheet.id}": str(ws.id),
            "${worksheet.status}": ws.status or "Open",
            "${worksheet.created_at}": _format_dt(ws.created_at),
            "${worksheet.breakdown_time}": _format_dt(ws.breakdown_time),
            "${worksheet.repair_finished_time}": _format_dt(ws.repair_finished_time),
            "${worksheet.downtime}": f"{ws.total_downtime_hours:.2f}h" if ws.total_downtime_hours is not None else "-",
            "${worksheet.assigned_username}": ws.assigned_user.full_name if ws.assigned_user else "-",
            "${worksheet.notes}": ws.notes or "-",
            "${machine.name}": ws.machine.name if ws.machine else "-",
            "${machine.serial_number}": ws.machine.serial_number if ws.machine and ws.machine.serial_number else "-",
            "${machine.manufacturer_model}": f"{ws.machine.manufacturer or '-'} / {ws.machine.model or '-'}" if ws.machine else "-",
            "${production_line.name}": ws.machine.production_line.name if ws.machine and ws.machine.production_line else "-",
            "${generated_at}": _format_dt(utcnow()),
            "${generated_by}": generated_by,
        }
        
        # Add parts information if available
        if ws.parts:
            parts_list = []
            for link in ws.parts:
                part = link.part
                parts_list.append(f"{part.sku if part else '-'} - {part.name if part else '-'} ({link.quantity_used}x)")
            replacements["${worksheet.parts}"] = "\n".join(parts_list) if parts_list else "-"
        else:
            replacements["${worksheet.parts}"] = "-"
        
        # Load template, replace placeholders, save as DOCX
        doc = _replace_placeholders_in_docx(template_path, replacements)
        
        # Populate parts table dynamically
        _populate_parts_table(doc, ws.parts if ws.parts else [])
        
        doc.save(str(docx_path))
        
        # Save record (storing DOCX path)
        pdf_record = session.query(WorksheetPDF).filter_by(worksheet_id=ws.id).first()
        if not pdf_record:
            pdf_record = WorksheetPDF(worksheet_id=ws.id)
        pdf_record.pdf_path = str(docx_path)  # Store DOCX path
        pdf_record.generated_at = utcnow()
        pdf_record.page_count = 1
        session.add(pdf_record)
        session.commit()

        logger.info(f"Worksheet DOCX generated: {docx_path}")
        return str(docx_path)
    finally:
        if should_close:
            session.close()


def generate_work_request_pdf(pm_task_id: int, generated_by: str = "system", output_path: str | None = None, session: Session = None) -> str:
    """Generate work request PDF for PM task (ISO 9001)"""
    session, should_close = _get_session(session)
    try:
        task: PMTask = (
            session.query(PMTask)
            .options(
                joinedload(PMTask.machine).joinedload(Machine.production_line),
                joinedload(PMTask.assigned_user),
                joinedload(PMTask.created_by_user),
            )
            .filter(PMTask.id == pm_task_id)
            .first()
        )
        if not task:
            raise ValueError("PM feladat nem található")

        if output_path:
            docx_path = Path(output_path)
            if docx_path.suffix.lower() != '.docx':
                docx_path = docx_path.with_suffix('.docx')
            docx_path.parent.mkdir(parents=True, exist_ok=True)
        else:
            output_dir = _ensure_output_dir()
            docx_path = output_dir / f"work_request_{task.id}.docx"

        # Use DOCX template if available (check selected template first)
        template_path = get_selected_work_request_template(session) or (TEMPLATES_DIR / "work_request_template.docx")
        if template_path.exists() and DOCX_AVAILABLE:
            # Prepare replacements
            replacements = {
                "${pm_task.id}": str(task.id),
                "${pm_task.task_name}": task.task_name or "-",
                "${pm_task.task_type}": task.task_type or "-",
                "${pm_task.priority}": task.priority or "-",
                "${pm_task.status}": task.status or "-",
                "${pm_task.due_date}": _format_dt(task.due_date or task.next_due_date),
                "${pm_task.estimated_duration_minutes}": f"{task.estimated_duration_minutes} perc / minutes" if task.estimated_duration_minutes else "-",
                "${pm_task.task_description}": task.task_description or "-",
                "${pm_task.location}": task.location or "-",
                "${pm_task.created_at}": _format_dt(task.created_at),
                "${machine.name}": task.machine.name if task.machine else "-",
                "${machine.serial_number}": task.machine.serial_number if task.machine and task.machine.serial_number else "-",
                "${machine.manufacturer_model}": f"{task.machine.manufacturer or '-'} / {task.machine.model or '-'}" if task.machine else "-",
                "${production_line.name}": task.machine.production_line.name if task.machine and task.machine.production_line else "-",
                "${assigned_user.full_name}": task.assigned_user.full_name if task.assigned_user else "Globális / Global",
                "${assignment_type}": "Személyre szabott / Assigned" if task.assigned_user else "Globális / Global",
                "${created_by_user.full_name}": task.created_by_user.full_name if task.created_by_user else generated_by,
                "${generated_at}": _format_dt(utcnow()),
                "${generated_by}": generated_by,
            }
            
            # Load template, replace placeholders, save as DOCX
            doc = _replace_placeholders_in_docx(template_path, replacements)
            doc.save(str(docx_path))
        else:
            # No fallback - template is required
            raise FileNotFoundError(f"Work request template not found: {template_path}")
        
        # Save record (storing DOCX path)
        pdf_record = session.query(WorkRequestPDF).filter_by(pm_task_id=task.id).first()
        if not pdf_record:
            pdf_record = WorkRequestPDF(pm_task_id=task.id)
        pdf_record.pdf_path = str(docx_path)  # Store DOCX path
        pdf_record.generated_at = utcnow()
        pdf_record.page_count = 1
        if task.created_by_user_id:
            pdf_record.generated_by_user_id = task.created_by_user_id
        session.add(pdf_record)
        session.commit()

        logger.info(f"Work request DOCX generated: {docx_path}")
        return str(docx_path)
    finally:
        if should_close:
            session.close()


def generate_pm_worksheet_pdf(pm_history_id: int, generated_by: str = "system", output_path: str | None = None, session: Session = None) -> str:
    """Generate worksheet PDF for completed PM task (ISO 9001)"""
    session, should_close = _get_session(session)
    try:
        history: PMHistory = (
            session.query(PMHistory)
            .options(
                joinedload(PMHistory.pm_task).joinedload(PMTask.machine).joinedload(Machine.production_line),
                joinedload(PMHistory.completed_user),
                joinedload(PMHistory.worksheet),
            )
            .filter(PMHistory.id == pm_history_id)
            .first()
        )
        if not history:
            raise ValueError("PM history nem található")
        
        task = history.pm_task

        if output_path:
            docx_path = Path(output_path)
            if docx_path.suffix.lower() != '.docx':
                docx_path = docx_path.with_suffix('.docx')
            docx_path.parent.mkdir(parents=True, exist_ok=True)
        else:
            output_dir = _ensure_output_dir()
            docx_path = output_dir / f"pm_worksheet_{history.id}.docx"

        # Use worksheet DOCX template (same as regular worksheets)
        template_path = get_selected_worksheet_template(session) or (TEMPLATES_DIR / "worksheet_template.docx")
        if template_path.exists() and DOCX_AVAILABLE:
            # Prepare replacements for worksheet template
            replacements = {
                "${worksheet.id}": str(history.id),
                "${worksheet.status}": history.completion_status or "completed",
                "${worksheet.created_at}": _format_dt(history.executed_date),
                "${worksheet.breakdown_time}": _format_dt(history.executed_date),
                "${worksheet.repair_finished_time}": _format_dt(history.executed_date),
                "${worksheet.downtime}": f"{history.duration_minutes or 0} perc / minutes",
                "${worksheet.assigned_username}": history.completed_user.full_name if history.completed_user else "-",
                "${worksheet.notes}": history.notes or "-",
                "${machine.name}": task.machine.name if task.machine else (task.location or "-"),
                "${machine.serial_number}": task.machine.serial_number if task.machine and task.machine.serial_number else "-",
                "${machine.manufacturer_model}": f"{task.machine.manufacturer or '-'} / {task.machine.model or '-'}" if task.machine else "-",
                "${production_line.name}": task.machine.production_line.name if task.machine and task.machine.production_line else "-",
                "${generated_at}": _format_dt(utcnow()),
                "${generated_by}": generated_by,
            }
            
            # Load template, replace placeholders, save as DOCX
            doc = _replace_placeholders_in_docx(template_path, replacements)
            
            # Add parts information if worksheet has parts
            if history.worksheet and history.worksheet.parts:
                parts_list = []
                for link in history.worksheet.parts:
                    part = link.part
                    parts_list.append(f"{part.sku if part else '-'} - {part.name if part else '-'} ({link.quantity_used}x)")
                # Try to find and replace parts placeholder in document
                for paragraph in doc.paragraphs:
                    if "${worksheet.parts}" in paragraph.text:
                        paragraph.text = paragraph.text.replace("${worksheet.parts}", "\n".join(parts_list) if parts_list else "-")
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if "${worksheet.parts}" in paragraph.text:
                                    paragraph.text = paragraph.text.replace("${worksheet.parts}", "\n".join(parts_list) if parts_list else "-")
            
            doc.save(str(docx_path))
        else:
            # No fallback - template is required
            raise FileNotFoundError(f"PM worksheet template not found: {template_path}")
        
        # Save record (storing DOCX path)
        pdf_record = session.query(PMWorksheetPDF).filter_by(pm_history_id=history.id).first()
        if not pdf_record:
            pdf_record = PMWorksheetPDF(pm_history_id=history.id)
        pdf_record.pdf_path = str(docx_path)  # Store DOCX path
        pdf_record.generated_at = utcnow()
        pdf_record.page_count = 1
        if history.completed_by_user_id:
            pdf_record.generated_by_user_id = history.completed_by_user_id
        session.add(pdf_record)
        session.commit()

        logger.info(f"PM worksheet DOCX generated: {docx_path}")
        return str(docx_path)
    finally:
        if should_close:
            session.close()


def merge_pdfs(pdf_paths: list[str], output_path: str) -> str:
    """Merge multiple PDFs into one document"""
    if not PYPDF2_AVAILABLE:
        raise ImportError("PyPDF2 or pypdf not available for PDF merging")
    
    if not pdf_paths:
        raise ValueError("No PDF paths provided")
    
    writer = PdfWriter()
    
    for pdf_path in pdf_paths:
        if not Path(pdf_path).exists():
            logger.warning(f"PDF not found: {pdf_path}")
            continue
        
        try:
            reader = PdfReader(pdf_path)
            for page in reader.pages:
                writer.add_page(page)
        except Exception as e:
            logger.error(f"Error reading PDF {pdf_path}: {e}")
            continue
    
    output_path_obj = Path(output_path)
    output_path_obj.parent.mkdir(parents=True, exist_ok=True)
    
    with open(output_path, 'wb') as output_file:
        writer.write(output_file)
    
    logger.info(f"Merged {len(pdf_paths)} PDFs into {output_path}")
    return output_path


def generate_scrapping_docx(
    entity_type: str,
    entity_id: int,
    reason: str,
    template_path: Optional[Path] = None,
    worksheet_id: Optional[int] = None,
    pm_history_id: Optional[int] = None,
    session: Session = None
) -> Path:
    """
    Generate scrapping DOCX document for a part or machine
    
    Args:
        entity_type: "Part" or "Machine"
        entity_id: ID of the part or machine
        reason: Reason for scrapping
        template_path: Optional template path (if None, uses selected template from settings)
        session: Database session
    
    Returns:
        Path: Path to generated DOCX file
    """
    session, should_close = _get_session(session)
    try:
        # Get entity
        if entity_type == "Part":
            entity = session.query(Part).filter_by(id=entity_id).first()
            if not entity:
                raise ValueError(f"Part with ID {entity_id} not found")
            entity_name = entity.name
            entity_identifier = entity.sku
            entity_type_display = "Alkatrész"
        elif entity_type == "Machine":
            entity = session.query(Machine).filter_by(id=entity_id).first()
            if not entity:
                raise ValueError(f"Machine with ID {entity_id} not found")
            entity_name = entity.name
            entity_identifier = entity.serial_number or entity.asset_tag or f"ID: {entity.id}"
            entity_type_display = "Gép"
        else:
            raise ValueError(f"Invalid entity_type: {entity_type}")
        
        # Get template path
        if template_path is None:
            template_path = get_selected_scrapping_template()
        
        if template_path is None or not template_path.exists():
            # Use default template or create a simple one
            template_path = TEMPLATES_DIR / "default_scrapping_template.docx"
            if not template_path.exists():
                # Create a simple default template
                if not DOCX_AVAILABLE:
                    raise ImportError("python-docx not available for DOCX generation")
                
                doc = Document()
                doc.add_heading('Selejtezési Lap / Scrapping Document', 0)
                doc.add_paragraph(f'Típus / Type: {entity_type_display}')
                doc.add_paragraph(f'Név / Name: {entity_name}')
                doc.add_paragraph(f'Azonosító / Identifier: {entity_identifier}')
                doc.add_paragraph(f'Selejtezés oka / Reason: {reason}')
                doc.add_paragraph(f'Dátum / Date: {_format_dt(utcnow())}')
                template_path.parent.mkdir(parents=True, exist_ok=True)
                doc.save(str(template_path))
        
        # Prepare output path
        output_dir = _ensure_output_dir()
        docx_path = output_dir / f"scrapping_{entity_type.lower()}_{entity_id}_{utcnow().strftime('%Y%m%d_%H%M%S')}.docx"
        
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available for DOCX generation")
        
        # Get current user
        from services.context_service import get_current_user_id, get_user
        user_id = get_current_user_id()
        user = get_user(user_id) if user_id else None
        user_name = user.full_name if user and user.full_name else (user.username if user else "System")
        
        # Get worksheet information if available
        worksheet_id_str = "-"
        worksheet_title = "-"
        if worksheet_id:
            from database.models import Worksheet
            worksheet = session.query(Worksheet).filter_by(id=worksheet_id).first()
            if worksheet:
                worksheet_id_str = str(worksheet.id)
                worksheet_title = worksheet.title or f"Munkalap #{worksheet.id}"
        
        # Get PM task information if available
        pm_task_name = "-"
        if pm_history_id:
            from database.models import PMHistory
            pm_history = session.query(PMHistory).options(
                joinedload(PMHistory.pm_task)
            ).filter_by(id=pm_history_id).first()
            if pm_history and pm_history.pm_task:
                pm_task_name = pm_history.pm_task.task_name or "-"
        
        # Build document dynamically with only fields that have values
        doc = Document()
        
        # Page setup
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)
        
        # Title
        title = doc.add_heading('Selejtezési Lap / Scrapping Document', 0)
        title_run = title.runs[0] if title.runs else title.add_run('Selejtezési Lap / Scrapping Document')
        title_run.font.color.rgb = RGBColor(0x00, 0x37, 0x6F)
        
        # Create data table with only fields that have values
        data_rows = []
        
        # Always include these fields
        data_rows.append(("Típus / Type", entity_type_display))
        data_rows.append(("Név / Name", entity_name))
        data_rows.append(("Azonosító / Identifier", entity_identifier))
        data_rows.append(("Selejtezés oka / Reason", reason))
        data_rows.append(("Selejtezte / Scrapped by", user_name))
        data_rows.append(("Selejtezés dátuma / Scrapped at", _format_dt(utcnow())))
        
        # Only include worksheet info if available
        if worksheet_id and worksheet_id_str != "-":
            data_rows.append(("Munkalap ID / Worksheet ID", worksheet_id_str))
            if worksheet_title and worksheet_title != "-":
                data_rows.append(("Munkalap címe / Worksheet Title", worksheet_title))
        
        # Only include PM task info if available
        if pm_history_id and pm_task_name and pm_task_name != "-":
            data_rows.append(("PM feladat neve / PM Task Name", pm_task_name))
        
        # Create table
        table = doc.add_table(rows=len(data_rows), cols=2)
        table.style = 'Light Grid Accent 1'
        table.autofit = False
        
        # Set column widths
        for row in table.rows:
            row.cells[0].width = Cm(6)
            row.cells[1].width = Cm(10)
        
        # Fill table with data
        for i, (label, value) in enumerate(data_rows):
            # Label cell (left column)
            label_cell = table.rows[i].cells[0]
            label_para = label_cell.paragraphs[0]
            label_run = label_para.add_run(label)
            label_run.bold = True
            label_run.font.size = Pt(11)
            label_run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
            if WD_ALIGN_VERTICAL:
                label_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            # Set background color for label cell
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'E7F3FF')
            label_cell._element.get_or_add_tcPr().append(shading_elm)
            
            # Value cell (right column)
            value_cell = table.rows[i].cells[1]
            value_para = value_cell.paragraphs[0]
            value_run = value_para.add_run(str(value))
            value_run.font.size = Pt(11)
            value_run.font.color.rgb = RGBColor(0x37, 0x47, 0x5A)
            if WD_ALIGN_VERTICAL:
                value_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        
        # Add notes section
        doc.add_paragraph()  # Spacer
        notes_heading = doc.add_heading('Megjegyzések / Notes', level=2)
        notes_para = doc.add_paragraph(
            'Ez a dokumentum automatikusan generálva lett a rendszer által.\n'
            'This document was automatically generated by the system.'
        )
        notes_para.paragraph_format.left_indent = Cm(0.5)
        notes_run = notes_para.runs[0] if notes_para.runs else notes_para.add_run(
            'Ez a dokumentum automatikusan generálva lett a rendszer által.\n'
            'This document was automatically generated by the system.'
        )
        notes_run.font.size = Pt(10)
        notes_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        notes_run.italic = True
        
        # Save document
        doc.save(str(docx_path))
        
        logger.info(f"Scrapping DOCX generated: {docx_path}")
        return docx_path
    finally:
        if should_close:
            session.close()


def generate_vacation_document(
    vacation_request_id: int,
    output_path: Optional[str] = None,
    session: Session = None
) -> Path:
    """
    Generate vacation request document (DOCX) for an approved vacation request
    
    Args:
        vacation_request_id: Vacation request ID
        output_path: Optional output path (if None, generates in default directory)
        session: Database session
    
    Returns:
        Path: Path to generated DOCX file
    """
    from services.settings_service import get_setting
    
    session, should_close = _get_session(session)
    
    try:
        # Get vacation request
        vacation_request = (
            session.query(VacationRequest)
            .options(
                joinedload(VacationRequest.user),
                joinedload(VacationRequest.approved_by),
            )
            .filter(VacationRequest.id == vacation_request_id)
            .first()
        )
        
        if not vacation_request:
            raise ValueError(f"Vacation request with ID {vacation_request_id} not found")
        
        # Get template path from settings
        template_path_str = get_setting("vacation_document_template", session=session)
        if template_path_str:
            template_path = Path(template_path_str)
        else:
            template_path = TEMPLATES_DIR / "default_vacation_template.docx"
        
        # Create default template if it doesn't exist
        if not template_path.exists():
            if not DOCX_AVAILABLE:
                raise ImportError("python-docx not available for DOCX generation")
            
            doc = Document()
            doc.add_heading('Szabadság Igénylő Lap / Vacation Request Form', 0)
            doc.add_paragraph('Igénylő neve / Applicant Name: ${user.full_name}')
            doc.add_paragraph('Felhasználónév / Username: ${user.username}')
            doc.add_paragraph('Kezdő dátum / Start Date: ${vacation.start_date}')
            doc.add_paragraph('Befejező dátum / End Date: ${vacation.end_date}')
            doc.add_paragraph('Napok száma / Days Count: ${vacation.days_count}')
            doc.add_paragraph('Indoklás / Reason: ${vacation.reason}')
            doc.add_paragraph('Igénylés dátuma / Requested At: ${vacation.requested_at}')
            doc.add_paragraph('Jóváhagyó neve / Approver Name: ${approver.full_name}')
            doc.add_paragraph('Jóváhagyás dátuma / Approved At: ${vacation.approved_at}')
            doc.add_paragraph('Generálás dátuma / Generated At: ${generated_at}')
            template_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(template_path))
        
        # Prepare output path
        if output_path:
            docx_path = Path(output_path)
            if docx_path.suffix.lower() != '.docx':
                docx_path = docx_path.with_suffix('.docx')
            docx_path.parent.mkdir(parents=True, exist_ok=True)
        else:
            output_dir = _ensure_output_dir()
            docx_path = output_dir / f"vacation_request_{vacation_request_id}_{utcnow().strftime('%Y%m%d_%H%M%S')}.docx"
        
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available for DOCX generation")
        
        # Get user information
        user = vacation_request.user
        approver = vacation_request.approved_by
        
        # Format dates
        start_date_str = vacation_request.start_date.strftime("%Y-%m-%d") if isinstance(vacation_request.start_date, datetime) else str(vacation_request.start_date.date())
        end_date_str = vacation_request.end_date.strftime("%Y-%m-%d") if isinstance(vacation_request.end_date, datetime) else str(vacation_request.end_date.date())
        requested_at_str = _format_dt(vacation_request.requested_at) if vacation_request.requested_at else "-"
        approved_at_str = _format_dt(vacation_request.approved_at) if vacation_request.approved_at else "-"
        
        # Prepare replacements
        replacements = {
            "${user.full_name}": user.full_name if user and user.full_name else (user.username if user else "-"),
            "${user.username}": user.username if user else "-",
            "${vacation.start_date}": start_date_str,
            "${vacation.end_date}": end_date_str,
            "${vacation.days_count}": str(vacation_request.days_count) if vacation_request.days_count else "-",
            "${vacation.reason}": vacation_request.reason or "-",
            "${vacation.requested_at}": requested_at_str,
            "${approver.full_name}": approver.full_name if approver and approver.full_name else (approver.username if approver else "-"),
            "${vacation.approved_at}": approved_at_str,
            "${generated_at}": _format_dt(utcnow()),
        }
        
        # Load template, replace placeholders, save as DOCX
        doc = _replace_placeholders_in_docx(template_path, replacements)
        doc.save(str(docx_path))
        
        # Save document record
        from services.context_service import get_current_user_id
        generated_by_user_id = get_current_user_id()
        
        vacation_doc = VacationDocument(
            vacation_request_id=vacation_request_id,
            docx_path=str(docx_path),
            generated_by_user_id=generated_by_user_id
        )
        session.add(vacation_doc)
        session.commit()
        
        logger.info(f"Vacation document generated: {docx_path}")
        return docx_path
        
    except Exception as e:
        session.rollback()
        logger.error(f"Error generating vacation document: {e}")
        raise
    finally:
        if should_close:
            session.close()
