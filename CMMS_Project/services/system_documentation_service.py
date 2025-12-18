"""
System Documentation Service
Teljes rendszer leírás generálása DOCX formátumban
"""

from pathlib import Path
from typing import Dict, List, Optional
from datetime import datetime
import logging
import re

logger = logging.getLogger(__name__)

# Try to import python-docx
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None
    logger.warning("python-docx not available, DOCX generation will not work")


def generate_system_documentation_docx(output_path: Optional[Path] = None) -> Path:
    """
    Generál egy teljes DOCX dokumentációt a rendszerről
    
    Args:
        output_path: Output fájl elérési út (opcionális)
    
    Returns:
        Path: A generált DOCX fájl elérési útja
    
    Raises:
        ImportError: Ha python-docx nincs telepítve
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is required for DOCX generation. Install it with: pip install python-docx")
    
    if output_path is None:
        output_dir = Path("generated_docs")
        output_dir.mkdir(exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = output_dir / f"system_documentation_{timestamp}.docx"
    
    doc = _build_docx_documentation()
    
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    
    logger.info(f"System documentation DOCX generated: {output_path}")
    return output_path


def _build_docx_documentation() -> Document:
    """Build the complete DOCX documentation"""
    
    # Create a new Document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Read documentation files
    # Use path relative to project root
    project_root = Path(__file__).parent.parent
    docs_dir = project_root / "docs"
    
    # Title page
    _add_title_page(doc)
    
    # Table of contents
    _add_table_of_contents(doc)
    
    # Add sections
    sections_data = [
        ("rendszer_attekintes", "Rendszer Áttekintés", docs_dir / "SYSTEM_ARCHITECTURE_ANALYSIS.md"),
        ("logikai_analizis", "Mélyreható Logikai Analízis", docs_dir / "DEEP_LOGICAL_ANALYSIS.md"),
        ("logikai_fa", "Logikai Fa Diagramok", docs_dir / "LOGICAL_TREE_DIAGRAM.md"),
        ("munkafolyamatok", "Munkafolyamatok", docs_dir / "SYSTEM_WORKFLOW_DIAGRAM.md"),
        ("entitasok_muveletei", "Entitás Műveletek", None),  # Will be built from code
        ("javaslatok", "Javaslatok és Javítások", docs_dir / "SYSTEM_IMPROVEMENTS_RECOMMENDATIONS.md"),
        ("osszefoglalo", "Összefoglaló", docs_dir / "COMPLETE_ANALYSIS_SUMMARY.md"),
    ]
    
    for section_key, section_title, doc_file in sections_data:
        if doc_file and doc_file.exists():
            content = _read_doc_file(doc_file)
            _add_section(doc, section_title, content)
        elif section_key == "entitasok_muveletei":
            # Build entity operations from code
            content = _build_entity_operations_content()
            _add_section(doc, section_title, content)
    
    return doc


def _read_doc_file(file_path: Path) -> Optional[str]:
    """Read a documentation file"""
    try:
        if file_path.exists():
            return file_path.read_text(encoding='utf-8')
    except Exception as e:
        logger.warning(f"Could not read documentation file {file_path}: {e}")
    return None


def _add_title_page(doc: Document):
    """Add title page to document"""
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("CMMS Rendszer")
    title_run.font.size = Pt(32)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run("Teljes Dokumentáció")
    subtitle_run.font.size = Pt(24)
    subtitle_run.font.color.rgb = RGBColor(64, 64, 64)
    
    doc.add_paragraph()  # Spacing
    
    version_para = doc.add_paragraph()
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_run = version_para.add_run("Karbantartás Menedzsment Rendszer v1.0.6")
    version_run.font.size = Pt(14)
    version_run.font.italic = True
    
    doc.add_paragraph()  # Spacing
    doc.add_paragraph()  # Spacing
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Generálva: {datetime.now().strftime('%Y.%m.%d %H:%M')}")
    date_run.font.size = Pt(12)
    
    # Page break
    doc.add_page_break()


def _add_table_of_contents(doc: Document):
    """Add table of contents"""
    toc_title = doc.add_heading("Tartalomjegyzék", 1)
    
    toc_items = [
        "Rendszer Áttekintés",
        "Mélyreható Logikai Analízis",
        "Logikai Fa Diagramok",
        "Munkafolyamatok",
        "Entitás Műveletek",
        "Javaslatok és Javítások",
        "Összefoglaló",
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Bullet')
        p_format = p.runs[0].font
        p_format.size = Pt(12)
    
    doc.add_page_break()


def _add_section(doc: Document, title: str, content: str):
    """Add a section to the document"""
    if not content:
        return
    
    # Section title
    heading = doc.add_heading(title, 1)
    
    # Parse and add markdown content
    _add_markdown_content(doc, content)


def _add_markdown_content(doc: Document, markdown_text: str):
    """Add markdown content to document"""
    if not markdown_text:
        return
    
    lines = markdown_text.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Empty line
        if not line:
            doc.add_paragraph()
            i += 1
            continue
        
        # Code blocks - add as formatted text
        if line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            
            if code_lines:
                # Add code block as monospace paragraph
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'No Spacing'
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                # Add spacing after code block
                doc.add_paragraph()
            continue
        
        # Headers
        if line.startswith('# '):
            heading = doc.add_heading(line[2:].strip(), 1)
            i += 1
            continue
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:].strip(), 2)
            i += 1
            continue
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:].strip(), 3)
            i += 1
            continue
        elif line.startswith('#### '):
            heading = doc.add_heading(line[5:].strip(), 4)
            i += 1
            continue
        
        # Lists
        if line.startswith('- ') or line.startswith('* '):
            list_items = []
            while i < len(lines) and (lines[i].strip().startswith('- ') or lines[i].strip().startswith('* ')):
                item_text = lines[i].strip()[2:].strip()
                list_items.append(item_text)
                i += 1
            
            for item in list_items:
                p = doc.add_paragraph(item, style='List Bullet')
                _format_paragraph_with_markdown(p, item)
            
            continue
        
        # Numbered lists
        numbered_match = re.match(r'^(\d+)\.\s+(.+)', line)
        if numbered_match:
            list_items = []
            while i < len(lines) and re.match(r'^\d+\.\s+', lines[i].strip()):
                item_text = re.sub(r'^\d+\.\s+', '', lines[i].strip())
                list_items.append(item_text)
                i += 1
            
            for item in list_items:
                p = doc.add_paragraph(style='List Number')
                _format_paragraph_with_markdown(p, item)
            
            continue
        
        # Regular paragraph
        # Check if it's part of a multi-line paragraph
        paragraph_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not _is_special_line(lines[i].strip()):
            paragraph_lines.append(lines[i].rstrip())
            i += 1
        
        paragraph_text = ' '.join(paragraph_lines)
        p = doc.add_paragraph()
        _format_paragraph_with_markdown(p, paragraph_text)


def _is_special_line(line: str) -> bool:
    """Check if line is a special markdown element"""
    return (line.startswith('#') or 
            line.startswith('- ') or 
            line.startswith('* ') or
            line.startswith('```') or
            re.match(r'^\d+\.\s+', line))


def _format_paragraph_with_markdown(p, text: str):
    """Add text to paragraph with markdown formatting (bold, italic)"""
    # Split by markdown markers and preserve formatting
    parts = re.split(r'(\*\*[^*]+\*\*|__[^_]+__|\*[^*]+\*|_[^_]+_|`[^`]+`)', text)
    
    for part in parts:
        if not part:
            continue
        
        # Bold: **text** or __text__
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('__') and part.endswith('__'):
            run = p.add_run(part[2:-2])
            run.bold = True
        # Italic: *text* or _text_
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = p.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('_') and part.endswith('_') and not part.startswith('__'):
            run = p.add_run(part[1:-1])
            run.italic = True
        # Code: `text`
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            run = p.add_run(part)
            run.font.size = Pt(11)


def _build_entity_operations_content() -> str:
    """Build entity operations step-by-step descriptions"""
    return """# Entitás Műveletek - Lépésről Lépésre Leírások

## PM Task (Preventive Maintenance Task)

### Létrehozás
1. PM menü → "Új feladat" gomb
2. Gép kiválasztása
3. Feladat neve és leírás megadása
4. Prioritás beállítása
5. Határidő beállítása
6. Hozzárendelés (globális vagy felhasználóhoz)
7. Mentés

**Automatikus folyamatok:**
- Workflow validáció (workflow_service)
- Notification küldés (notification_service)
- Work Request PDF generálás (pdf_service)
- SystemLog bejegyzés (log_service)

### Elvégzés
1. PM menü → "Elvégzés" gomb a feladathoz
2. Completion dialog megnyitása
3. Dátum és idő megadása
4. Munka leírása
5. Megfigyelések rögzítése
6. Státusz választás (kész/részleges/problémás)
7. Alkatrészek hozzáadása (opcionális):
   - Part kiválasztás
   - Mennyiség megadása
   - Storage location kiválasztás
8. Fájlok feltöltése (opcionális)
9. Mentés

**Automatikus folyamatok:**
- PMHistory létrehozása
- Worksheet automatikus létrehozása (ha alkatrész használva)
- WorkRequestPDF generálás
- PMWorksheetPDF generálás
- Scrapping Documents generálás (ha alkatrész)
- Fájlok mentése (PMTaskAttachment)
- Notification küldés
- SystemLog bejegyzés

## Worksheet (Munkalap)

### Létrehozás
1. Worksheets menü → "Új munkalap" gomb
2. Gép kiválasztása
3. Cím és leírás megadása
4. Hozzárendelt felhasználó kiválasztása
5. Mentés

**Automatikus folyamatok:**
- Notification küldés
- SystemLog bejegyzés

### Alkatrész hozzáadása
1. Worksheet részletek → "Alkatrész hozzáadása" gomb
2. Part kiválasztása
3. Mennyiség megadása
4. Storage location kiválasztása
5. Mentés

**Automatikus folyamatok:**
- Készlet ellenőrzés (InventoryLevel)
- PartLocation frissítés (storage_service)
- StockTransaction létrehozás (transaction_service)
- Scrapping Document generálás (ha auto-generate enabled)
- SystemLog bejegyzés

### Lezárás
1. Worksheet részletek → "Lezárás" gomb
2. Javítás befejezés időpontjának megadása
3. Mentés

**Automatikus folyamatok:**
- Workflow validáció
- Downtime kalkuláció
- WorksheetPDF generálás
- Scrapping Documents generálás (ha van alkatrész)
- Notification küldés
- SystemLog bejegyzés

## Part (Alkatrész)

### Létrehozás
1. Inventory/Alkatrészek menü → "Új alkatrész" gomb
2. SKU megadása (validáció: egyedi, formátum ellenőrzés)
3. Név megadása
4. Kategória, beszállító (opcionális)
5. Ár információk (vételár, eladási ár)
6. Készlet szint információk (safety stock, reorder quantity)
7. Kezdeti mennyiség (opcionális)
8. Storage location hozzárendelés (opcionális)
9. Mentés

**Automatikus folyamatok:**
- SKU validáció (validators.validate_sku)
- InventoryLevel automatikus létrehozás (quantity_on_hand=0)
- StockTransaction létrehozás (ha initial_quantity > 0)
- PartLocation létrehozás (ha storage_location megadva)
- SystemLog bejegyzés

### Készletbevétel
1. Inventory/Alkatrészek menü → Part kiválasztás → "Készletbevétel" gomb
2. Mennyiség megadása
3. Storage location kiválasztása (csak releváns helyek: üres vagy azonos SKU)
4. Mentés

**Automatikus folyamatok:**
- PartLocation létrehozás/frissítés
- InventoryLevel frissítés
- StockTransaction létrehozás
- Validáció (InventoryLevel ↔ PartLocation)
- SystemLog bejegyzés

## Machine (Gép)

### Létrehozás
1. Assets menü → Production Line kiválasztás → "Új gép" gomb
2. Alap információk: név, sorozatszám, modell, gyártó
3. Dátum információk: install date, purchase date, warranty expiry
4. Operációs információk: operating hours, maintenance interval
5. Fizikai információk: tömeg, méretek
6. Finanszírozási információk: purchase price, supplier
7. Kompatibilis alkatrészek hozzárendelése
8. Mentés

**Automatikus folyamatok:**
- SystemLog bejegyzés
- AssetHistory bejegyzés

### Karbantartás igénylése
1. Production Line menü → Gép kiválasztás → "Karbantartás igénylése" gomb
2. Feladat neve és leírás megadása
3. Hozzárendelés (globális vagy felhasználóhoz)
4. Prioritás beállítása
5. Határidő beállítása
6. Mentés

**Automatikus folyamatok:**
- PMTask létrehozása (pm_service.create_pm_task)
- Notification küldés
- Work Request PDF generálás

## Storage Location (Raktárhely)

### Létrehozás
1. Storage menü → "Új raktárhely" gomb
2. Név megadása
3. Szülő hely kiválasztása (opcionális - hierarchikus struktúra)
4. Típus, kód, leírás (opcionális)
5. Mentés

**Automatikus folyamatok:**
- Hierarchia validáció (cirkuláris referencia ellenőrzés)
- SystemLog bejegyzés

### Alkatrész hozzárendelés
1. Storage menü → Part kiválasztás → "Tárhelyhez rendelés" gomb
2. Storage location kiválasztása (csak releváns: üres vagy azonos SKU)
3. Mennyiség megadása
4. Mentés

**Automatikus folyamatok:**
- PartLocation létrehozás/frissítés
- InventoryLevel frissítés
- Validáció (InventoryLevel ↔ PartLocation)
- SystemLog bejegyzés
"""
