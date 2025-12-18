#!/usr/bin/env python3
"""Test paragraph replacement logic"""

# Simulate the problematic function
def test_paragraph_replacement():
    from docx import Document
    from docx.oxml.ns import qn
    
    # Create a test document
    doc = Document()
    p = doc.add_paragraph("Hello ${name}, you have ${count} items")
    
    print(f"Before: {p.text}")
    print(f"Runs: {len(p.runs)}")
    
    # Test the replacement logic
    original_text = p.text
    substituted_text = original_text.replace("${name}", "John").replace("${count}", "5")
    
    print(f"Substituted: {substituted_text}")
    
    if original_text == substituted_text:
        print("No change needed")
        return
    
    # Clear all runs and rebuild
    for _ in range(len(p.runs)):
        run_elem = p.runs[0]._element
        run_elem.getparent().remove(run_elem)
    
    print(f"Runs after clear: {len(p.runs)}")
    
    # Add new run
    p.add_run(substituted_text)
    
    print(f"After: {p.text}")
    print(f"Runs: {len(p.runs)}")

if __name__ == '__main__':
    try:
        test_paragraph_replacement()
        print("SUCCESS")
    except Exception as e:
        print(f"ERROR: {e}")
        import traceback
        traceback.print_exc()
