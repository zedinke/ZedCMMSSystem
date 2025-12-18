"""
QR Code Generation Utility
Generates QR codes for parts and creates printable labels
"""

import qrcode
from qrcode.image.pil import PilImage
from PIL import Image, ImageDraw, ImageFont
from pathlib import Path
from typing import List, Optional
from io import BytesIO
import base64

from database.models import Part, QRCodeData
from database.session_manager import SessionLocal
from sqlalchemy.orm import Session
from datetime import datetime

import logging

logger = logging.getLogger(__name__)


def generate_qr_code(part_id: int, sku: str, size: int = 200) -> Optional[PilImage]:
    """
    Generate QR code image for a part
    
    Args:
        part_id: Part ID
        sku: Part SKU code
        size: Image size in pixels (default: 200)
    
    Returns:
        PIL Image object with QR code, or None on error
    """
    try:
        # Create QR code instance
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=10,
            border=4,
        )
        
        # QR data: part_id:sku format for easy scanning
        qr_data = f"{part_id}:{sku}"
        qr.add_data(qr_data)
        qr.make(fit=True)
        
        # Create image
        img = qr.make_image(fill_color="black", back_color="white")
        
        # Resize if needed
        if size != 200:
            img = img.resize((size, size), Image.Resampling.LANCZOS)
        
        logger.info(f"QR code generated for part {part_id} (SKU: {sku})")
        return img
        
    except Exception as e:
        logger.error(f"Error generating QR code for part {part_id}: {e}")
        return None


def save_qr_code_to_file(part_id: int, sku: str, output_path: Path, size: int = 200) -> bool:
    """
    Generate and save QR code to file
    
    Args:
        part_id: Part ID
        sku: Part SKU code
        output_path: Path to save the image
        size: Image size in pixels
    
    Returns:
        True if successful, False otherwise
    """
    try:
        img = generate_qr_code(part_id, sku, size)
        if img is None:
            return False
        
        # Ensure directory exists
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Save image
        img.save(output_path, format='PNG')
        
        logger.info(f"QR code saved to {output_path}")
        return True
        
    except Exception as e:
        logger.error(f"Error saving QR code to {output_path}: {e}")
        return False


def generate_qr_code_base64(part_id: int, sku: str, size: int = 200) -> Optional[str]:
    """
    Generate QR code and return as base64 string
    
    Args:
        part_id: Part ID
        sku: Part SKU code
        size: Image size in pixels
    
    Returns:
        Base64 encoded image string, or None on error
    """
    try:
        img = generate_qr_code(part_id, sku, size)
        if img is None:
            return None
        
        # Convert to base64
        buffer = BytesIO()
        img.save(buffer, format='PNG')
        img_str = base64.b64encode(buffer.getvalue()).decode()
        
        return f"data:image/png;base64,{img_str}"
        
    except Exception as e:
        logger.error(f"Error generating QR code base64 for part {part_id}: {e}")
        return None


def generate_qr_labels(part_ids: List[int], output_format: str = 'docx', 
                      label_size: tuple = (100, 50), session: Session = None, 
                      user_id: int = None, output_path: Optional[Path] = None) -> Optional[Path]:
    """
    Generate printable QR code labels for multiple parts in DOCX format
    
    Args:
        part_ids: List of part IDs
        output_format: 'docx' (default: 'docx') - PDF is no longer supported
        label_size: Label size in mm (width, height) - for reference only
        session: Database session
        user_id: ID of user generating the labels (optional, will try to get from context if not provided)
    
    Returns:
        Path to generated DOCX file, or None on error
    """
    try:
        from docx import Document
        from docx.shared import Cm, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        logger.error("python-docx not installed. Please install it with: pip install python-docx")
        return None
    
    from database.models import User
    
    session, should_close = _get_session(session)
    
    try:
        # Get current user if not provided
        user_info = None
        if user_id is None:
            try:
                from services.context_service import get_current_user
                current_user = get_current_user()
                if current_user:
                    user_info = current_user.full_name or current_user.username
            except:
                pass
        elif user_id:
            user = session.query(User).filter(User.id == user_id).first()
            if user:
                user_info = user.full_name or user.username
        
        # Get parts from database with relationships
        parts = session.query(Part).filter(Part.id.in_(part_ids)).all()
        
        if not parts:
            logger.warning("No parts found for QR label generation")
            return None
        
        # Use provided output path or create default
        if output_path is None:
            # Create output directory
            output_dir = Path("data/reports/qr_labels")
            output_dir.mkdir(parents=True, exist_ok=True)
            
            # Generate DOCX with labels
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = output_dir / f"qr_labels_{timestamp}.docx"
        else:
            # Ensure parent directory exists
            output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Check if template is selected
        from services.settings_service import get_selected_qr_label_template
        from config.app_config import TEMPLATES_DIR
        template_path = get_selected_qr_label_template(session)
        
        # Create DOCX document (from template if available, otherwise new)
        if template_path and template_path.exists():
            doc = Document(str(template_path))
            logger.info(f"Using QR label template: {template_path}")
        else:
            doc = Document()
            # Set default page margins if no template
            sections = doc.sections
            for section in sections:
                section.top_margin = Cm(1)
                section.bottom_margin = Cm(1)
                section.left_margin = Cm(1)
                section.right_margin = Cm(1)
        
        # Calculate labels per row (A4 width: ~21cm, margins: 2cm, label width: ~10cm)
        labels_per_row = 2
        
        # Generate QR codes and create labels
        row_labels = []
        
        for i, part in enumerate(parts):
            # Generate QR code image
            qr_img = generate_qr_code(part.id, part.sku, size=100)
            if qr_img is None:
                continue
            
            # Save QR code temporarily (use output_path's parent directory)
            temp_qr_path = output_path.parent / f"temp_qr_{part.id}.png"
            qr_img.save(temp_qr_path)
            
            # Get machine information for this part
            compatible_machines = part.compatible_machines if hasattr(part, 'compatible_machines') else []
            parent_info = ""
            child_info = ""
            
            if compatible_machines:
                # Get first compatible machine as parent
                parent_machine = compatible_machines[0]
                if parent_machine:
                    # Try to get production line as parent
                    if parent_machine.production_line:
                        parent_info = parent_machine.production_line.name
                    else:
                        parent_info = parent_machine.name
                    
                    # Try to get module as child (if exists)
                    if hasattr(parent_machine, 'modules') and parent_machine.modules:
                        child_machine = parent_machine.modules[0]
                        if child_machine:
                            child_info = child_machine.name
            
            # Get creation date
            creation_date = part.created_at.strftime("%Y-%m-%d %H:%M") if part.created_at else datetime.now().strftime("%Y-%m-%d %H:%M")
            
            # Store label data for this part
            label_data = {
                'sku': part.sku,
                'name': part.name,
                'parent_info': parent_info,
                'child_info': child_info,
                'user_info': user_info,
                'creation_date': creation_date,
                'qr_path': temp_qr_path
            }
            
            row_labels.append(label_data)
            
            # When we have enough labels per row, or it's the last part, create a row table
            if (i + 1) % labels_per_row == 0 or i == len(parts) - 1:
                # Create a horizontal table for this row
                row_table = doc.add_table(rows=1, cols=len(row_labels))
                row_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                for j, label_data in enumerate(row_labels):
                    cell = row_table.cell(0, j)
                    cell.width = Cm(label_size[0] / 10)  # Convert mm to cm
                    
                    # Create a 2-column table inside the cell: left = data, right = QR code
                    inner_table = cell.add_table(rows=1, cols=2)
                    inner_table.width = Cm(label_size[0] / 10)
                    
                    # Left cell: text data
                    left_cell = inner_table.cell(0, 0)
                    left_cell.width = Cm((label_size[0] / 10) * 0.6)  # 60% for text
                    p_left = left_cell.paragraphs[0]
                    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # SKU
                    run = p_left.add_run(f"SKU: {label_data['sku']}\n")
                    run.font.size = Pt(9)
                    run.font.bold = True
                    
                    # Name
                    run = p_left.add_run(f"{label_data['name'][:15]}\n")
                    run.font.size = Pt(8)
                    
                    # Machine information
                    if label_data['parent_info']:
                        run = p_left.add_run(f"Szülő: {label_data['parent_info'][:12]}\n")
                        run.font.size = Pt(7)
                    if label_data['child_info']:
                        run = p_left.add_run(f"Alárendelt: {label_data['child_info'][:10]}\n")
                        run.font.size = Pt(7)
                    
                    # Creator and date
                    if label_data['user_info']:
                        run = p_left.add_run(f"Készítette:\n{label_data['user_info'][:10]}\n")
                        run.font.size = Pt(7)
                    run = p_left.add_run(f"Dátum:\n{label_data['creation_date']}")
                    run.font.size = Pt(7)
                    
                    # Right cell: QR code
                    right_cell = inner_table.cell(0, 1)
                    right_cell.width = Cm((label_size[0] / 10) * 0.4)  # 40% for QR code
                    p_right = right_cell.paragraphs[0]
                    p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Add QR code image
                    p_right.add_run().add_picture(str(label_data['qr_path']), width=Cm(2.5))
                
                row_labels = []
                # Add spacing between rows
                doc.add_paragraph()
        
        # Save document
        doc.save(str(output_path))
        
        # Clean up temp files
        for temp_file in output_path.parent.glob("temp_qr_*.png"):
            temp_file.unlink()
        
        logger.info(f"QR labels DOCX generated: {output_path}")
        return output_path
            
    except Exception as e:
        logger.error(f"Error generating QR labels: {e}")
        import traceback
        traceback.print_exc()
        return None
    finally:
        if should_close:
            session.close()


def record_qr_code_generation(part_id: int, qr_data: str, session: Session = None) -> Optional[QRCodeData]:
    """
    Record QR code generation in database
    
    Args:
        part_id: Part ID
        qr_data: QR code data string
        session: Database session
    
    Returns:
        QRCodeData object, or None on error
    """
    session, should_close = _get_session(session)
    
    try:
        qr_record = QRCodeData(
            part_id=part_id,
            qr_data=qr_data,
            generated_at=datetime.now()
        )
        
        session.add(qr_record)
        session.commit()
        session.refresh(qr_record)
        
        logger.info(f"QR code generation recorded for part {part_id}")
        return qr_record
        
    except Exception as e:
        logger.error(f"Error recording QR code generation for part {part_id}: {e}")
        session.rollback()
        return None
    finally:
        if should_close:
            session.close()


def _get_session(session: Session = None) -> tuple:
    """
    Get database session, creating one if needed
    
    Returns:
        Tuple of (session, should_close) where should_close indicates if session was created here
    """
    if session is None:
        return SessionLocal(), True
    return session, False
